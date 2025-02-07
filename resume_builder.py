import json

from docx import Document
from docx.shared import Pt, RGBColor

def set_font_style(paragraph, font_name="Arial", font_size=12, bold=False, color=(0, 0, 0)):
    """Sets font styling for a paragraph."""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.font.color.rgb = RGBColor(*color)

def add_section(doc, title):
    """Adds a section title to the document."""
    para = doc.add_paragraph(title)
    set_font_style(para, font_size=14, bold=True, color=(0, 0, 255))

def add_personal_info(doc, personal_info):
    header = doc.add_paragraph()
    run = header.add_run(f"{personal_info['full_name']} | {personal_info['email']} | {personal_info['phone']}")
    set_font_style(header, font_size=14, bold=True)
    header = doc.add_paragraph()
    run = header.add_run(f"{personal_info['address']} | {personal_info['linkedin']} | {personal_info['website']}")
    set_font_style(header, font_size=12, color=(50, 50, 50))

def add_summary(doc, summary):
    """Adds summary section to the document."""
    add_section(doc, "Summary")
    doc.add_paragraph(summary)

def add_work_experience(doc, work_experience):
    """Adds work experience section."""
    add_section(doc, "Work Experience")
    for job in work_experience:
        para = doc.add_paragraph(f"{job['position']} at {job['company']} ({job['location']})")
        set_font_style(para, font_size=12, bold=True)
        date_range = f"{job['start_date']} - {job['end_date'] if job['end_date'] else 'Present'}"
        doc.add_paragraph(date_range)
        for desc in job["description"]:
            doc.add_paragraph(f"{desc}", style="ListBullet")

def add_education(doc, education):
    """Adds education section."""
    add_section(doc, "Education")
    for edu in education:
        para = doc.add_paragraph(f"{edu['degree']} in {edu['field_of_study']}, {edu['institution']}")
        set_font_style(para, font_size=12, bold=True)
        date_range = f"{edu['start_date']} - {edu['end_date']}"
        doc.add_paragraph(date_range)
        doc.add_paragraph(f"GPA: {edu['grade']}")
        
def add_skills(doc, skills):
    """Adds skills section."""
    add_section(doc, "Skills")
    for skill in skills:
        doc.add_paragraph(f"{skill}", style="ListBullet")

def convert_cv_json_to_docx(cv_json, output_filename):
    """Converts CV JSON to a well-formatted DOCX file."""
    doc = Document()
    add_personal_info(doc, cv_json['personal_info'])
    add_summary(doc, cv_json['summary'])
    add_work_experience(doc, cv_json['work_experience'])
    add_education(doc, cv_json['education'])
    add_skills(doc, cv_json['skills'])
    doc.save(output_filename)
    print(f"CV saved as {output_filename}")

# Example Usage
cv_json = {
    "personal_info": {
        "full_name": "John Doe",
        "email": "johndoe@example.com",
        "phone": "+1-555-1234-567",
        "address": "123 Main Street, Springfield, IL, USA",
        "linkedin": "https://www.linkedin.com/in/johndoe",
        "website": "https://johndoe.dev"
    },
    "summary": "Experienced AWS Solutions Architect with 10+ years of expertise in cloud migrations, security, and infrastructure automation. Passionate about helping businesses scale efficiently using modern cloud technologies.",
    "work_experience": [
        {
        "company": "Argorand Inc.",
        "position": "Founder & Principal Consultant",
        "location": "Remote",
        "start_date": "2024-01-01",
        "end_date": None,
        "description": [
                "Leading a cloud consulting firm specializing in AWS migration and implementation services for FinTech SMEs and federal government clients.",
                "Leading a cloud consulting firm specializing in AWS migration and implementation services for FinTech SMEs and federal government clients.",
                "Leading a cloud consulting firm specializing in AWS migration and implementation services for FinTech SMEs and federal government clients."
        ]
        },
        {
        "company": "Ciklum",
        "position": "AWS Solutions Architect",
        "location": "New York, NY",
        "start_date": "2022-03-01",
        "end_date": "2023-12-31",
        "description": [
            "Designed and implemented AWS cloud solutions for enterprise clients, optimizing performance and reducing costs by 30%.",
            "Designed and implemented AWS cloud solutions for enterprise clients, optimizing performance and reducing costs by 30%.",
            "Designed and implemented AWS cloud solutions for enterprise clients, optimizing performance and reducing costs by 30%.",
        ]
        },
    {
      "company": "Satsyil Corp",
      "position": "Cloud Engineer",
      "location": "Washington, DC",
      "start_date": "2019-06-01",
      "end_date": "2022-02-28",
      "description": [
        "Managed cloud migrations for federal agencies, ensuring compliance with FedRAMP and security best practices."
      ]
    }
    ],
  "education": [
        {
            "institution": "University of Illinois at Urbana-Champaign",
            "degree": "Master of Science",
            "field_of_study": "Computer Science",
            "start_date": "2017-08-01",
            "end_date": "2019-05-30",
            "grade": "3.9/4.0"
        },
        {
            "institution": "University of Illinois at Urbana-Champaign",
            "degree": "Bachelor of Science",
            "field_of_study": "Information Technology",
            "start_date": "2013-08-01",
            "end_date": "2017-05-30",
            "grade": "3.8/4.0"
        }
    ],   
  
  "skills": [
    "AWS Cloud Architecture",
    "Infrastructure as Code (IaC)",
    "Terraform",
    "Kubernetes",
    "Serverless Computing",
    "Cloud Security",
    "DevOps",
    "Python",
    "CI/CD Pipelines"
  ]
}

convert_cv_json_to_docx(cv_json, "John_Doe_CV.docx")

